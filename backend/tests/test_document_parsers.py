"""
Tests for Module 6's document parsers.

Generates REAL PDF/DOCX/TXT files on the fly (using the same libraries
that write them: fitz for PDF, python-docx for DOCX) and verifies our
extraction round-trips the content correctly. This is a genuine
end-to-end test, not a mock.

EasyOCR (image parsing) is tested separately -- see the note in
test_ocr_parser.py regarding sandbox verification status.

Run with: pytest backend/tests/test_document_parsers.py -v
"""

import pathlib

import fitz
import pytest
from docx import Document

from app.utils.parsers.document_parser import (
    EmptyDocumentError,
    UnsupportedFileTypeError,
    extract_text,
)
from app.utils.parsers.docx_parser import extract_text_from_docx
from app.utils.parsers.pdf_parser import extract_text_from_pdf
from app.utils.parsers.txt_parser import extract_text_from_txt

SAMPLE_TEXT = "This is a sample essay about renewable energy and climate policy."


@pytest.fixture
def sample_pdf(tmp_path) -> str:
    """Generate a real PDF with actual selectable text (not scanned/image)."""
    path = tmp_path / "sample.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), SAMPLE_TEXT)
    doc.save(str(path))
    doc.close()
    return str(path)


@pytest.fixture
def sample_docx(tmp_path) -> str:
    """Generate a real DOCX with two paragraphs."""
    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph(SAMPLE_TEXT)
    doc.add_paragraph("A second paragraph about solar panel adoption.")
    doc.save(str(path))
    return str(path)


@pytest.fixture
def sample_txt(tmp_path) -> str:
    path = tmp_path / "sample.txt"
    path.write_text(SAMPLE_TEXT, encoding="utf-8")
    return str(path)


@pytest.fixture
def empty_txt(tmp_path) -> str:
    path = tmp_path / "empty.txt"
    path.write_text("   \n\n  ", encoding="utf-8")
    return str(path)


def test_pdf_extraction(sample_pdf):
    text = extract_text_from_pdf(sample_pdf)
    assert "renewable energy" in text


def test_docx_extraction(sample_docx):
    text = extract_text_from_docx(sample_docx)
    assert "renewable energy" in text
    assert "solar panel" in text


def test_txt_extraction(sample_txt):
    text = extract_text_from_txt(sample_txt)
    assert text == SAMPLE_TEXT


def test_txt_handles_cp1252_encoding(tmp_path):
    path = tmp_path / "windows_encoded.txt"
    # U+2019 (right single quotation mark, "'") is a common character in
    # Word-exported text that encodes to byte 0x92 in cp1252 but has no
    # representation in strict ASCII/UTF-8-assuming-cp1252 readers.
    path.write_bytes("It\u2019s a test".encode("cp1252"))
    text = extract_text_from_txt(str(path))
    assert "test" in text  # doesn't crash, extracts what it can


def test_facade_dispatches_pdf(sample_pdf):
    text = extract_text(sample_pdf, "essay.pdf")
    assert "renewable energy" in text


def test_facade_dispatches_docx(sample_docx):
    text = extract_text(sample_docx, "essay.docx")
    assert "solar panel" in text


def test_facade_dispatches_txt(sample_txt):
    text = extract_text(sample_txt, "essay.txt")
    assert text == SAMPLE_TEXT


def test_facade_rejects_unsupported_extension(sample_txt):
    with pytest.raises(UnsupportedFileTypeError):
        extract_text(sample_txt, "essay.xyz")


def test_facade_rejects_empty_txt(empty_txt):
    with pytest.raises(EmptyDocumentError):
        extract_text(empty_txt, "empty.txt")


def test_facade_flags_scanned_pdf_with_no_text_layer(tmp_path):
    """A blank PDF page (no text inserted) simulates a scanned document."""
    path = tmp_path / "scanned.pdf"
    doc = fitz.open()
    doc.new_page()  # blank page, no text layer
    doc.save(str(path))
    doc.close()

    with pytest.raises(EmptyDocumentError, match="scanned"):
        extract_text(str(path), "scanned.pdf")
